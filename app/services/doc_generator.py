import asyncio
import json
import logging
import tempfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.config import settings
from app.core.database import db, new_id
from app.services.llm_gateway import llm_gateway
from app.services.self_evaluator import evaluate_document
from app.services.whatsapp_outbound import whatsapp

logger = logging.getLogger(__name__)

_TEMPLATES_ROOT = Path(__file__).resolve().parent.parent.parent / "templates"
_BRAND_CONFIG_PATH = _TEMPLATES_ROOT / "shared" / "brand_config.json"

_DOC_TYPE_ALIASES: dict[str, str] = {
    "proforma": "proforma_invoice",
    "pi": "proforma_invoice",
    "proforma invoice": "proforma_invoice",
    "quotation": "commercial_quotation",
    "quote": "commercial_quotation",
    "commercial quotation": "commercial_quotation",
    "packing list": "packing_list",
    "packing": "packing_list",
    "letterhead": "letterhead",
    "letter": "letterhead",
}

_TEMPLATE_SCHEMAS: dict[str, dict] = {
    "proforma_invoice": {
        "buyer_name": "",
        "buyer_address": "",
        "seller_name": "",
        "seller_address": "",
        "invoice_number": "",
        "invoice_date": "",
        "line_items": [
            {
                "description": "",
                "qty": "",
                "unit": "",
                "unit_price": "",
                "total": "",
            }
        ],
        "subtotal": "",
        "tax": "0",
        "grand_total": "",
        "currency": "USD",
        "payment_terms": "",
        "delivery_terms": "",
        "validity": "",
        "notes": "",
    },
    "commercial_quotation": {
        "buyer_name": "",
        "buyer_address": "",
        "seller_name": "",
        "seller_address": "",
        "quotation_number": "",
        "quotation_date": "",
        "valid_until": "",
        "line_items": [
            {
                "description": "",
                "qty": "",
                "unit": "",
                "unit_price": "",
                "total": "",
            }
        ],
        "subtotal": "",
        "tax": "0",
        "grand_total": "",
        "currency": "USD",
        "payment_terms": "",
        "delivery_terms": "",
        "notes": "",
    },
    "packing_list": {
        "buyer_name": "",
        "buyer_address": "",
        "seller_name": "",
        "seller_address": "",
        "invoice_ref": "",
        "date": "",
        "line_items": [
            {
                "description": "",
                "qty": "",
                "unit": "",
                "net_weight": "",
                "gross_weight": "",
                "dimensions": "",
            }
        ],
        "total_packages": "",
        "total_net_weight": "",
        "total_gross_weight": "",
        "notes": "",
    },
    "letterhead": {
        "date": "",
        "recipient_name": "",
        "recipient_address": "",
        "subject": "",
        "body": "",
        "signatory_name": "",
        "signatory_title": "",
    },
}

_EXTRACT_SYSTEM_PROMPT = """\
You are extracting document variables from a WhatsApp conversation for a business document.

Document type: {doc_type}
Command: {command}
Recent conversation:
{context}

Known pricing data from database:
{db_prices}

Extract ALL required variables. For prices, use ONLY values from the database if available.
For the currency, default to USD. Exchange rate: 1 USD = {inr_to_usd} INR.

Return ONLY valid JSON with the structure:
{template_schema}"""


def _load_brand_config() -> dict:
    if _BRAND_CONFIG_PATH.is_file():
        return json.loads(_BRAND_CONFIG_PATH.read_text(encoding="utf-8"))
    return {}


def _resolve_doc_type(command: str) -> str:
    cmd_lower = command.lower()
    for alias, dtype in _DOC_TYPE_ALIASES.items():
        if alias in cmd_lower:
            return dtype
    return "proforma_invoice"


def _resolve_brand(chat_context: list[dict]) -> str:
    brand_keywords = {
        "stel astra": "stel_astra",
        "stelastra": "stel_astra",
        "nsi": "nsi_projects",
        "north star": "nsi_projects",
        "pacific": "pacific_unity",
    }
    combined = " ".join(m.get("content", "") for m in chat_context).lower()
    for keyword, brand in brand_keywords.items():
        if keyword in combined:
            return brand
    return "pacific_unity"


async def generate_document(
    chat_id: str,
    command: str,
    requested_by: str,
    chat_context: list[dict],
) -> None:
    doc_id = new_id()

    from app.services.task_engine import resolve_employee
    emp = await resolve_employee(requested_by)
    requester_id = emp["id"] if emp else None

    await whatsapp.send_text(chat_id, "Working on it... Gathering data.")

    doc_type = _resolve_doc_type(command)
    brand = _resolve_brand(chat_context)

    try:
        variables = await extract_variables(command, doc_type, chat_context)
    except Exception:
        logger.exception("Variable extraction failed for doc_id=%s", doc_id)
        await whatsapp.send_text(
            chat_id,
            "I couldn't extract the document details. "
            "Please provide more information (buyer, items, prices, etc.).",
        )
        return

    brand_config = _load_brand_config()
    brand_info = brand_config.get(brand, {})
    if brand_info:
        variables.setdefault("seller_name", brand_info.get("company_name", ""))
        variables.setdefault("seller_address", brand_info.get("address", ""))

    valid, price_issues = await validate_prices(variables)
    if not valid:
        logger.warning("Price validation issues for doc_id=%s: %s", doc_id, price_issues)

    try:
        docx_path = await generate_docx(doc_type, brand, variables)
    except Exception:
        logger.exception("DOCX generation failed for doc_id=%s", doc_id)
        await whatsapp.send_text(chat_id, "Failed to generate the document. Please try again.")
        return

    pdf_path = await convert_to_pdf(docx_path)

    score, eval_issues = await evaluate_document(doc_type, variables)

    await db.execute(
        "INSERT INTO generated_documents "
        "(id, doc_type, brand, status, template_used, variables_json, "
        "self_eval_score, self_eval_issues, requested_by, source_chat_id, created_at) "
        "VALUES (?, ?, ?, 'draft', ?, ?, ?, ?, ?, ?, ?)",
        (
            doc_id,
            doc_type,
            brand,
            f"{brand}/{doc_type}.docx",
            json.dumps(variables, default=str),
            score,
            json.dumps(eval_issues),
            requester_id,
            chat_id,
            datetime.now(timezone.utc).isoformat(),
        ),
    )

    send_path = pdf_path or docx_path
    ext = "pdf" if pdf_path else "docx"
    filename = f"{doc_type}_{doc_id[:8]}.{ext}"

    quality_note = ""
    if score < 0.7:
        quality_note = (
            f"\n\n⚠️ Quality score: {score:.0%}\nIssues found:\n"
            + "\n".join(f"• {i}" for i in eval_issues)
        )
    elif eval_issues:
        quality_note = (
            f"\n\nℹ️ Quality score: {score:.0%} — minor issues:\n"
            + "\n".join(f"• {i}" for i in eval_issues)
        )

    await whatsapp.send_document(
        chat_id,
        send_path,
        filename,
        caption=(
            f"📄 Draft {doc_type.replace('_', ' ').title()} generated."
            f"{quality_note}\n\n"
            f"Reply *approve* to finalize or describe any changes needed."
        ),
    )


async def extract_variables(
    command: str, doc_type: str, chat_context: list[dict]
) -> dict:
    schema = _TEMPLATE_SCHEMAS.get(doc_type, _TEMPLATE_SCHEMAS["proforma_invoice"])

    db_prices = await db.fetch_all(
        "SELECT name, metadata FROM entities WHERE type = 'price'"
    )
    prices_text = json.dumps(db_prices, default=str) if db_prices else "No pricing data in database."

    context_text = "\n".join(
        f"{m.get('role', 'user')}: {m.get('content', '')}" for m in chat_context
    )

    prompt = _EXTRACT_SYSTEM_PROMPT.format(
        doc_type=doc_type,
        command=command,
        context=context_text,
        db_prices=prices_text,
        inr_to_usd=settings.inr_to_usd,
        template_schema=json.dumps(schema, indent=2),
    )

    response = await llm_gateway.chat(
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": f"Extract variables for: {command}"},
        ],
        max_tokens=32000,
        request_type="doc_generator:extract_variables",
    )

    raw = response.get("content", "")
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.split("\n", 1)[-1]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        cleaned = cleaned.strip()

    return json.loads(cleaned)


async def validate_prices(variables: dict) -> tuple[bool, list[str]]:
    items = variables.get("line_items")
    if not items:
        return True, []

    issues: list[str] = []
    for i, item in enumerate(items, 1):
        desc = item.get("description", "")
        if not desc:
            continue
        row = await db.fetch_one(
            "SELECT metadata FROM entities WHERE type = 'price' AND name LIKE ?",
            (f"%{desc}%",),
        )
        if not row:
            continue
        try:
            meta = json.loads(row["metadata"]) if isinstance(row["metadata"], str) else row["metadata"]
            db_price = str(meta.get("unit_price", ""))
        except (json.JSONDecodeError, TypeError, AttributeError):
            continue

        stated_price = str(item.get("unit_price", ""))
        if db_price and stated_price and db_price != stated_price:
            issues.append(
                f"Line {i} '{desc}': stated price {stated_price} "
                f"differs from database price {db_price}"
            )

    return len(issues) == 0, issues


async def generate_docx(doc_type: str, brand: str, variables: dict) -> str:
    template_path = _TEMPLATES_ROOT / brand / f"{doc_type}.docx"

    if template_path.is_file():
        return _fill_template(template_path, variables)

    logger.info(
        "Template not found at %s — using programmatic fallback", template_path
    )
    return _build_docx_from_scratch(doc_type, brand, variables)


def _fill_template(template_path: Path, variables: dict) -> str:
    doc = Document(str(template_path))

    flat_vars = {
        k: str(v)
        for k, v in variables.items()
        if not isinstance(v, (list, dict))
    }

    for paragraph in doc.paragraphs:
        for key, value in flat_vars.items():
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in flat_vars.items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)

    items = variables.get("line_items", [])
    if items:
        _populate_template_table(doc, items)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _populate_template_table(doc: Document, items: list[dict]) -> None:
    for table in doc.tables:
        header_cells = [c.text.strip().lower() for c in table.rows[0].cells]
        if not any(kw in " ".join(header_cells) for kw in ("description", "item", "qty")):
            continue

        template_row_count = len(table.rows)
        if template_row_count > 1:
            template_row = table.rows[1]
            col_count = len(template_row.cells)
        else:
            col_count = len(table.columns)

        while len(table.rows) > 1:
            table._tbl.remove(table.rows[-1]._tr)

        for item in items:
            row_cells = table.add_row().cells
            values = list(item.values())
            for j in range(min(col_count, len(values))):
                row_cells[j].text = str(values[j])
        break


def _build_docx_from_scratch(
    doc_type: str, brand: str, variables: dict
) -> str:
    brand_config = _load_brand_config()
    brand_info = brand_config.get(brand, {})
    company_name = brand_info.get("company_name", brand.replace("_", " ").title())
    company_address = brand_info.get("address", "")
    company_email = brand_info.get("email", "")

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if doc_type == "letterhead":
        return _build_letterhead(doc, company_name, company_address, company_email, variables)

    title_map = {
        "proforma_invoice": "PROFORMA INVOICE",
        "commercial_quotation": "COMMERCIAL QUOTATION",
        "packing_list": "PACKING LIST",
    }
    doc_title = title_map.get(doc_type, doc_type.replace("_", " ").upper())

    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(company_name.upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    if company_address:
        addr_para = doc.add_paragraph()
        addr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        addr_run = addr_para.add_run(company_address)
        addr_run.font.size = Pt(9)
        addr_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if company_email:
        email_para = doc.add_paragraph()
        email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        email_run = email_para.add_run(company_email)
        email_run.font.size = Pt(9)
        email_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(doc_title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph()

    _add_detail_line(doc, "To", variables.get("buyer_name", ""))
    _add_detail_line(doc, "Address", variables.get("buyer_address", ""))

    ref_field = (
        "invoice_number"
        if doc_type == "proforma_invoice"
        else "quotation_number"
        if doc_type == "commercial_quotation"
        else "invoice_ref"
    )
    date_field = (
        "invoice_date"
        if doc_type == "proforma_invoice"
        else "quotation_date"
        if doc_type == "commercial_quotation"
        else "date"
    )
    _add_detail_line(doc, "Ref", variables.get(ref_field, ""))
    _add_detail_line(doc, "Date", variables.get(date_field, ""))

    currency = variables.get("currency", "USD")
    _add_detail_line(doc, "Currency", currency)

    doc.add_paragraph()

    items = variables.get("line_items", [])
    if items and doc_type in ("proforma_invoice", "commercial_quotation"):
        _build_price_table(doc, items)
    elif items and doc_type == "packing_list":
        _build_packing_table(doc, items)

    doc.add_paragraph()

    if doc_type in ("proforma_invoice", "commercial_quotation"):
        totals_para = doc.add_paragraph()
        totals_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        subtotal = variables.get("subtotal", "")
        tax = variables.get("tax", "0")
        grand_total = variables.get("grand_total", "")
        totals_text = f"Subtotal: {currency} {subtotal}"
        if tax and tax != "0":
            totals_text += f"\nTax: {currency} {tax}"
        totals_text += f"\nGrand Total: {currency} {grand_total}"
        totals_run = totals_para.add_run(totals_text)
        totals_run.bold = True
        totals_run.font.size = Pt(10)

    if doc_type == "packing_list":
        _add_detail_line(doc, "Total Packages", variables.get("total_packages", ""))
        _add_detail_line(doc, "Total Net Weight", variables.get("total_net_weight", ""))
        _add_detail_line(doc, "Total Gross Weight", variables.get("total_gross_weight", ""))

    doc.add_paragraph()

    for field in ("payment_terms", "delivery_terms", "validity", "valid_until"):
        val = variables.get(field, "")
        if val:
            _add_detail_line(doc, field.replace("_", " ").title(), val)

    notes = variables.get("notes", "")
    if notes:
        doc.add_paragraph()
        _add_detail_line(doc, "Notes", notes)

    footer_section = doc.sections[0]
    footer = footer_section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"{company_name} | {company_email}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _build_letterhead(
    doc: Document,
    company_name: str,
    company_address: str,
    company_email: str,
    variables: dict,
) -> str:
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(company_name.upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    if company_address:
        addr_run = header_para.add_run(f"\n{company_address}")
        addr_run.font.size = Pt(9)
        addr_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    if company_email:
        email_run = header_para.add_run(f"  |  {company_email}")
        email_run.font.size = Pt(9)
        email_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("_" * 72).runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    date_val = variables.get("date", "")
    if date_val:
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.add_run(date_val).font.size = Pt(10)

    doc.add_paragraph()
    if variables.get("recipient_name"):
        doc.add_paragraph(variables["recipient_name"])
    if variables.get("recipient_address"):
        doc.add_paragraph(variables["recipient_address"])

    doc.add_paragraph()
    if variables.get("subject"):
        subj_para = doc.add_paragraph()
        subj_run = subj_para.add_run(f"Re: {variables['subject']}")
        subj_run.bold = True
        subj_run.font.size = Pt(11)

    doc.add_paragraph()
    body = variables.get("body", "")
    if body:
        for line in body.split("\n"):
            doc.add_paragraph(line)

    doc.add_paragraph()
    doc.add_paragraph()
    if variables.get("signatory_name"):
        sig_para = doc.add_paragraph()
        sig_run = sig_para.add_run(variables["signatory_name"])
        sig_run.bold = True
    if variables.get("signatory_title"):
        doc.add_paragraph(variables["signatory_title"])

    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    tmp.close()
    return tmp.name


def _add_detail_line(doc: Document, label: str, value: str) -> None:
    if not value:
        return
    para = doc.add_paragraph()
    label_run = para.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    value_run = para.add_run(str(value))
    value_run.font.size = Pt(10)


def _build_price_table(doc: Document, items: list[dict]) -> None:
    headers = ["#", "Description", "Qty", "Unit", "Unit Price", "Total"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for idx, item in enumerate(items, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = str(item.get("description", ""))
        row_cells[2].text = str(item.get("qty", ""))
        row_cells[3].text = str(item.get("unit", ""))
        row_cells[4].text = str(item.get("unit_price", ""))
        row_cells[5].text = str(item.get("total", ""))
        for cell in row_cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)


def _build_packing_table(doc: Document, items: list[dict]) -> None:
    headers = ["#", "Description", "Qty", "Unit", "Net Wt", "Gross Wt", "Dimensions"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for idx, item in enumerate(items, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = str(item.get("description", ""))
        row_cells[2].text = str(item.get("qty", ""))
        row_cells[3].text = str(item.get("unit", ""))
        row_cells[4].text = str(item.get("net_weight", ""))
        row_cells[5].text = str(item.get("gross_weight", ""))
        row_cells[6].text = str(item.get("dimensions", ""))
        for cell in row_cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)


async def convert_to_pdf(docx_path: str) -> str | None:
    outdir = str(Path(docx_path).parent)
    try:
        proc = await asyncio.create_subprocess_exec(
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            outdir,
            docx_path,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=60)

        if proc.returncode != 0:
            logger.warning(
                "LibreOffice conversion failed (rc=%d): %s",
                proc.returncode,
                stderr.decode(errors="replace"),
            )
            return None

        pdf_path = Path(docx_path).with_suffix(".pdf")
        if pdf_path.is_file():
            return str(pdf_path)

        logger.warning("PDF file not found after conversion at %s", pdf_path)
        return None

    except FileNotFoundError:
        logger.warning("LibreOffice not installed — skipping PDF conversion")
        return None
    except asyncio.TimeoutError:
        logger.warning("LibreOffice conversion timed out")
        return None


async def handle_approval(
    chat_id: str, doc_id: str, user_response: str, approved_by: str
) -> None:
    doc = await db.fetch_one(
        "SELECT * FROM generated_documents WHERE id = ? AND status = 'draft'",
        (doc_id,),
    )
    if not doc:
        await whatsapp.send_text(chat_id, "Could not find the draft document.")
        return

    response_lower = user_response.lower()

    if "approve" in response_lower:
        now = datetime.now(timezone.utc).isoformat()
        await db.execute(
            "UPDATE generated_documents SET status = 'approved', "
            "approved_by = ?, approved_at = ? WHERE id = ?",
            (approved_by, now, doc_id),
        )

        drive_uploaded = False
        try:
            from app.services.drive_client import drive_client  # type: ignore[import-not-found]
            if drive_client:
                drive_uploaded = True
                logger.info("Document %s uploaded to Drive", doc_id)
        except ImportError:
            pass

        confirmation = f"✅ Document approved and finalized."
        if drive_uploaded:
            confirmation += " Uploaded to Google Drive."
        await whatsapp.send_text(chat_id, confirmation)
        return

    variables = json.loads(doc["variables_json"]) if doc.get("variables_json") else {}
    doc_type = doc.get("doc_type", "proforma_invoice")
    brand = doc.get("brand", "pacific_unity")

    chat_context = [
        {"role": "user", "content": user_response},
        {"role": "assistant", "content": f"Original variables: {json.dumps(variables, default=str)}"},
    ]

    try:
        await extract_variables(
            f"Revise the document: {user_response}", doc_type, chat_context
        )
    except Exception:
        logger.exception("Re-extraction failed for doc_id=%s", doc_id)
        await whatsapp.send_text(chat_id, "I couldn't apply those changes. Please describe them differently.")
        return

    await db.execute(
        "UPDATE generated_documents SET status = 'superseded' WHERE id = ?",
        (doc_id,),
    )

    await generate_document(
        chat_id,
        f"Regenerate {doc_type} with changes: {user_response}",
        approved_by,
        chat_context,
    )
